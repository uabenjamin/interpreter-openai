from __future__ import annotations

import asyncio
import re
from dataclasses import dataclass
from pathlib import Path

from openai import OpenAI

from .error_handling import UserFacingError


SUPPORTED_REFERENCE_EXTENSIONS = {
    ".docx",
    ".markdown",
    ".md",
    ".pdf",
    ".rtf",
    ".text",
    ".txt",
}
MAX_RAW_REFERENCE_CHARS = 12_000
MAX_SUMMARY_INPUT_CHARS = 60_000
SUMMARY_MAX_OUTPUT_TOKENS = 1200


@dataclass(slots=True)
class ReferenceDocument:
    source_name: str
    text: str


def extract_reference_document(path: Path) -> ReferenceDocument:
    resolved = path.expanduser()
    if not resolved.exists():
        raise UserFacingError(f"Reference file not found: {resolved}")
    if not resolved.is_file():
        raise UserFacingError(f"Reference path is not a file: {resolved}")

    extension = resolved.suffix.casefold()
    if extension not in SUPPORTED_REFERENCE_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_REFERENCE_EXTENSIONS))
        raise UserFacingError(
            f"Unsupported reference file type '{extension}'. Supported: {supported}"
        )

    if extension == ".docx":
        text = _extract_docx_text(resolved)
    elif extension == ".pdf":
        text = _extract_pdf_text(resolved)
    elif extension == ".rtf":
        text = _extract_rtf_text(resolved)
    else:
        text = _extract_plain_text(resolved)

    normalized = _normalize_reference_text(text)
    if not normalized:
        raise UserFacingError(f"No readable text found in reference file: {resolved}")
    return ReferenceDocument(source_name=resolved.name, text=normalized)


def build_raw_reference_pack(source_name: str, raw_text: str) -> str:
    normalized = _normalize_reference_text(raw_text)
    if not normalized:
        raise UserFacingError("Reference text is empty.")

    truncated = normalized[:MAX_RAW_REFERENCE_CHARS]
    omitted = len(normalized) - len(truncated)
    note = (
        f"\n\nNote: raw reference was truncated; {omitted:,} characters were omitted."
        if omitted > 0
        else ""
    )
    return (
        "# Sermon Reference\n"
        f"Source: {source_name}\n"
        "Mode: raw excerpt\n\n"
        "Use this context to improve translation terminology, Bible passage wording, "
        "names, sermon flow, and ASR correction. Do not quote or summarize it unless "
        "the live English segment says it.\n\n"
        f"{truncated}{note}"
    )


async def summarize_reference_pack(
    client: OpenAI,
    model: str,
    target_language_label: str,
    source_name: str,
    raw_text: str,
) -> str:
    return await asyncio.to_thread(
        _summarize_reference_pack_blocking,
        client,
        model,
        target_language_label,
        source_name,
        raw_text,
    )


def _summarize_reference_pack_blocking(
    client: OpenAI,
    model: str,
    target_language_label: str,
    source_name: str,
    raw_text: str,
) -> str:
    normalized = _normalize_reference_text(raw_text)
    if not normalized:
        raise UserFacingError("Reference text is empty.")

    input_text = normalized[:MAX_SUMMARY_INPUT_CHARS]
    omitted = len(normalized) - len(input_text)
    response = client.responses.create(
        model=model,
        instructions=(
            "Create a compact sermon reference pack for a live church interpreter. "
            "The pack will be reused during English-to-target-language translation. "
            "Extract only useful translation context: sermon title/topic, Scripture "
            "references and likely Bible quotes, people/church/ministry names, "
            "announcements, key theological terms with preferred translations, "
            "sermon outline, recurring metaphors, and ASR correction hints. "
            "Do not write a devotional summary. Use concise Markdown bullets."
        ),
        input=(
            f"Target language: {target_language_label}\n"
            f"Source file: {source_name}\n"
            f"Input was truncated before summarization: {'yes' if omitted > 0 else 'no'}\n\n"
            "SERMON DRAFT / REFERENCE TEXT:\n"
            f"{input_text}"
        ),
        max_output_tokens=SUMMARY_MAX_OUTPUT_TOKENS,
        temperature=0.1,
    )
    summary = (response.output_text or "").strip()
    if not summary:
        raise RuntimeError("OpenAI returned an empty sermon reference summary.")

    truncation_note = (
        f"\n\nNote: summarization input was truncated; {omitted:,} characters were omitted."
        if omitted > 0
        else ""
    )
    return (
        "# Sermon Reference\n"
        f"Source: {source_name}\n"
        "Mode: summarized reference pack\n\n"
        "Use this context to improve translation terminology, Bible passage wording, "
        "names, sermon flow, and ASR correction. Do not quote or summarize it unless "
        "the live English segment says it.\n\n"
        f"{summary}{truncation_note}"
    )


def _extract_plain_text(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "utf-16", "mac_roman"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise UserFacingError(
            "Reading .docx files requires python-docx. Run `.venv/bin/pip install -e .`."
        ) from exc

    document = Document(str(path))
    chunks: list[str] = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))
    return "\n".join(chunks)


def _extract_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise UserFacingError(
            "Reading .pdf files requires pypdf. Run `.venv/bin/pip install -e .`."
        ) from exc

    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n\n".join(pages)


def _extract_rtf_text(path: Path) -> str:
    raw = _extract_plain_text(path)
    raw = re.sub(r"{\\[^{}]+}", " ", raw)
    raw = re.sub(r"\\'[0-9a-fA-F]{2}", " ", raw)
    raw = re.sub(r"\\[a-zA-Z]+\d* ?", " ", raw)
    raw = raw.replace("{", " ").replace("}", " ")
    return raw


def _normalize_reference_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
